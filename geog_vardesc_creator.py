#DHS Geograhpy Var Desc generator

#Written by Gretchen Corcoran

#Virtual environment is var_desc_geog

#for single sample var_descs, only for now

#this takes in a list of samples
#could also do it by looping through file, like variables.xlsx, and checking if relevant geog vardescs exist, in future, so don't need any manual interaction
"""
geog_vardesc_creator.py

Iterates through a list of single sample geography variable descriptions to generate, and outputs these according to a standard formatting.

Usage python3 
"""
import docx
import argparse
import os
import sys
import traceback
import pandas as pd
import re
import textwrap
from docx.shared import Pt
from docx.enum.text import WD_COLOR_INDEX


#geography list parser function
def _geog_list_reader(sample_list_txt_file: str) -> list[str]:
    """
    Reads list of sample names from provided text file. Changes to lowercase for better processing.

    Params: sample_list_txt_file (str): Path to a text file containig sample names. Default is /pkg/dhs/variables/geography/templates/vardescs_to_generate.txt

    Returns:
        list[str]: List of sample names like "AF2015" etc.
    """
    with open(sample_list_txt_file, "r") as sample_list:
        return [line.strip().lower() for line in sample_list if line.strip() and not line.strip().startswith("#")] #dont necessarily need to strip here probably. also skips ### lines
    
#function to pull country name
def _sample_country_tuple_getter(sample_list_txt_file: str) -> list[tuple[str, str]]:
    """
    Creates a list of tuple pairs, each with the sample code and the country name, e.g., (sn2023, Senegal)

    Params: sample_list_text_File (str): Path to a text file containing a list of DHS samples. Default is /pkg/dhs/variables/geography/templates/vardescs_to_generate.txt

    Returns:
        list[tuple[str, str]]: A list of tuple pairs, containing the sample abbreviation and the country name.
    """


    sample_list = _geog_list_reader(sample_list_txt_file) #returns list[str]

    countries_control_file_path = "/pkg/ipums/dhs/metadata/countries.xlsx"
    countries_df = pd.read_excel(countries_control_file_path, header=0, engine = "openpyxl")

    sample_country_tuple_list = []

    for sample in sample_list:
        #grab the name of the country in countries_df and put into tuple
        country_name = countries_df.loc[countries_df['country'] == sample[0:2], 'fullname'].iloc[0] #grabbing first two letters of sample, then its full name
        country_tuple = (sample, country_name)
        sample_country_tuple_list.append(country_tuple)
    return sample_country_tuple_list

#function to grab other geo files if applicable
def _find_additional_geo_vars(sample_abbreviation: str) -> list[str]:
    """
    Checks if additional multi-year geography variables exist for the country, and if so makes a list of them.

    Params: A string containing the sample abbreivation, e.g., "sn" - with only the country part of it.

    Returns: A list of strings, each the name of a multi-year DHS geography variable, e.g., "geo_sn1992_2019"
    """
    found_variable_files = []
    geo_string = f'geo_{sample_abbreviation.lower()}'
    year_pattern = re.compile(r"\d{4}_\d{4}")

    folder_path = '/pkg/ipums/dhs/variables/geography'

    for entry in os.listdir(folder_path):
        full_path = os.path.join(folder_path, entry)
        if os.path.isfile(full_path) and entry.lower().endswith(".xlsx"): #if is .xlsx, only count once not twice because has .doc and .xlsx files in folder
            if geo_string in entry.lower() and year_pattern.search(entry) and not 'delete' in entry.lower() and not '$' in entry.lower():
                found_variable_files.append(entry)

    return found_variable_files

#create xml_snippet
def _build_xml_snippet(sample_name, country_name, year, multiple_years = False, crvar_multi_years = None):
    """
    Creates one of two different xml snippets creating the metadata text for DHS single-sample geography variable description. One snippet is if multiple year geography variables exist, the other is if it doesn't (which is the case if there is only one year for the country).
    
    Params: sample_name (string), country_name (string), year of the sample (string), whether multiple-year samples exist (boolean), and crvar_multi_years, a list of multiple year geography variables.

    Returns: XML snippet in f string form.
    """
    if multiple_years:
        snippet = f"""<vardesc>

<var>
GEO_{sample_name}
</var>

<desc>
GEO_{sample_name} (V101_{sample_name}) indicates the region of {country_name} where the respondent was interviewed. DHS regions in the {year} {country_name} survey are equivalent to regions.

Other samples have their own sample-specific geography variables. There are other integrated variables that provide spatially consistent regions over time. These include {crvar_multi_years}

A GIS map for GEO_{sample_name} (in shapefile format) can be downloaded from the DHS program <link id="24">Spatial Data Repository</link> Boundaries page.
</desc>

<comp>
GEO_{sample_name} (V101_{sample_name}) is a country- and sample-specific variable and has no comparability issues.

<em>Comparability - Standard DHS</em>

GEO_{sample_name}, like other V101 variables, is a geographic variable added during processing of the DHS data. V101 is included in all Phases of the DHS.
</comp>

<comment>
</comment>

</vardesc>
"""
    else:
        snippet = f"""<vardesc>

<var>
GEO_{sample_name}
</var>

<desc>
GEO_{sample_name} (V101_{sample_name}) indicates the region of {country_name} where the respondent was interviewed. DHS regions in the {year} {country_name} survey are equivalent to regions.

A GIS map for GEO_{sample_name} (in shapefile format) can be downloaded from the DHS program <link id="24">Spatial Data Repository</link> Boundaries page.
</desc>

<comp>
GEO_{sample_name} (V101_{sample_name}) is a country- and sample-specific variable and has no comparability issues.

<em>Comparability - Standard DHS</em>

GEO_{sample_name}, like other V101 variables, is a geographic variable added during processing of the DHS data. V101 is included in all Phases of the DHS.
</comp>

<comment>
</comment>

</vardesc>
"""
    return textwrap.dedent(snippet).strip()

#function to create description
def _geog_vardesc_document_creator(sample_list_text_file: str) -> list[str]:
    """
    Creates a .doc for each variable 

    NOTE: I think it is fine that we are creating .docx, I believe it doesn't matter anymore. But should double check

    Params: A text link to a .txt list of researcher-specified samples

    Returns:
        list[str]: A list of geography variable descriptions created. Created variables are in /pkg/ipums/dhs/variables/geography/templates/vardescs_to_generate.txt
    """

    sample_country_tuple_list = _sample_country_tuple_getter(sample_list_text_file) #returns list[tuple]
    new_vardescs_created = []

    for sample_tuple in sample_country_tuple_list:
        crvar_multi_years = None #resetting
        if sample_tuple[0][-1].isdigit(): #if is in format AF2015
            sample_name = sample_tuple[0].upper() #uppercase if not
        else: #if is in format AF2015IR
            sample_name = sample_tuple[0][:7].upper() #takes first 6 digits, so first 2 letters + year
        country_name = sample_tuple[1]
        #year not just string slicing because what if enter as "bd2022ir" instead of "bd2022"
        year = re.search(r"\d{4}", sample_tuple[0]).group()

        #add wording alternate if multiple samples
        other_geog_list = _find_additional_geo_vars(sample_tuple[0][:2]) #finding multiple year geography
        cleaned_vars = [var.replace(".xlsx", "").rsplit("_", 1)[0] for var in other_geog_list]

        if cleaned_vars:
            multiple_years = True
            if len(cleaned_vars) == 1:
                crvar_multi_years = f'<crvar>{cleaned_vars[0]}</crvar>.'
            elif len(cleaned_vars) == 2:
                crvar_multi_years = f"<crvar>{cleaned_vars[0]}</crvar> and <crvar>{cleaned_vars[1]}</crvar>."
            else:
                crvar_multi_years = ", ".join(f"<crvar>{v}</crvar>" for v in cleaned_vars[:-1]) + f", and <crvar>{cleaned_vars[-1]}</crvar>."
        else:
            multiple_years = False
            crvar_multi_years = None
        # if sum(os.path.isdir(os.path.join(f'/pkg/ipums/dhs/country/{country_name.lower()}', entry))
        #        for entry in os.listdir(f'/pkg/ipums/dhs/country/{country_name.lower()}')) > 0:
        #     multiple_years = True
        #     other_geog_list = _find_additional_geo_vars(sample_tuple[0][:2])

        #     #remove suffixes from other geog_vars list
        #     cleaned_vars = [var.replace(".xlsx", "").rsplit("_", 1)[0] for var in other_geog_list]

        #     #crvars for multiple year samples, putting into sentence structure
        #     if len(cleaned_vars) == 1:
        #         crvar_multi_vars = f'<crvar>{cleaned_vars[0]}</crvar>.'
        #     elif len(cleaned_vars) == 2: #if is only two, put in a sentence
        #         crvar_multi_years = f"<crvar>{cleaned_vars[0]}</crvar> and <crvar>{cleaned_vars[1]}</crvar>."
        #     else: #if 3 or more. won't ever be 1 because of logic
        #         crvar_multi_years = ", ".join(f"<crvar>{v}</crvar>" for v in cleaned_vars[:-1]) + f", and <crvar>{cleaned_vars[-1]}</crvar>."
        # else:
        #     multiple_years = False

        #XML string for output - doesn't include bolded "Comparability - Standard DHS" line, hopefully is fine
        #currently not included "other sample years have their own sample-specific geography variables"


        xml_snippet = _build_xml_snippet(sample_name, country_name, year, multiple_years, crvar_multi_years)
        #output
        doc = docx.Document()

        lines = xml_snippet.strip().splitlines()

        #with crvar highlighting - chatgpt assisted this part
        crvar_pattern = re.compile(r"(<crvar>.*?</crvar>)")
        for line in lines:
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.space_before = Pt(0)

            cursor = 0
            matches = list(crvar_pattern.finditer(line))

            if matches:
                for match in matches:
                    start, end = match.span()
                    if start > cursor:
                        before_text = line[cursor:start]
                        run = paragraph.add_run(before_text)
                        run.font.name = "Times New Roman"
                    
                    crvar_text = match.group(1)
                    highlight_run = paragraph.add_run(crvar_text)
                    highlight_run.font.name = "Times New Roman"
                    highlight_run.font.highlight_color = WD_COLOR_INDEX.YELLOW

                    cursor = end
                if cursor < len(line):
                    after_text = line[cursor:]
                    run = paragraph.add_run(after_text)
                    run.font.name = "Times New Roman"
            else:
                run = paragraph.add_run(line)
                run.font.name = "Times New Roman"

        #this is output without crvar highlighting
        # for line in lines:
        #     paragraph = doc.add_paragraph()
        #     paragraph.paragraph_format.space_after = Pt(0)
        #     paragraph.paragraph_format.space_before = Pt(0)
        #     #paragraph.paragraph_format.line_spacing = 1
        #     run = paragraph.add_run(line)
        #     run.font.name = "Times New Roman"

        output_dir = '/pkg/ipums/dhs/variables/geography/templates/autogenerated_vardescs/'
        var_name_for_doc = f'geo_{sample_name.lower()}_desc.docx'
        new_geog_vardesc = os.path.join(output_dir, var_name_for_doc)
        new_vardescs_created.append(new_geog_vardesc)
        doc.save(new_geog_vardesc)

    return new_vardescs_created

#this has a default, auto reading a vardescs_to_generate but can be overwritten
def main():
    parser = argparse.ArgumentParser(
        description = "Creates geography variable descriptions based on a list of provided samples"
    )

    parser.add_argument("--sample_list", type=str, default = "/pkg/ipums/dhs/variables/geography/templates/vardescs_to_generate.txt")

    args = parser.parse_args()

    sample_list = args.sample_list

    if not os.path.exists(sample_list):
        print(f"Error: Wrong provided file path, you can always leave this part of the command out and use the default txt file here: /pkg/ipums/dhs/variables/geography/templates/vardescs_to_generate.txt ")
        sys.exit(1)

    sample_list = args.sample_list

    try:
        new_geog_vardescs_created = _geog_vardesc_document_creator(sample_list)

        print(f'\nFinished. Created {len(new_geog_vardescs_created)} single-sample geography variable descriptions, located in /pkg/ipums/dhs/variables/geography/templates/autogenerated_vardescs.')
    except Exception as e:
        print(f'An error occurred:', file = sys.stderr)
        print(str(e), file = sys.stderr)
        traceback.print_exc()
        sys.exit(2)

if __name__ == "__main__":
    main()